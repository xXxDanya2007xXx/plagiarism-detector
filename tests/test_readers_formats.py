from pathlib import Path

from plagiarism_detector.readers import read_document, read_folder


def test_read_docx(tmp_path: Path):
    import docx  # type: ignore

    p = tmp_path / "sample.docx"
    d = docx.Document()
    d.add_paragraph("Hello DOCX world")
    d.save(p)

    doc = read_document(p)
    assert "hello" in doc.text.lower()


def test_read_pdf_no_crash(tmp_path: Path):
    from pypdf import PdfWriter  # type: ignore

    p = tmp_path / "sample.pdf"
    w = PdfWriter()
    w.add_blank_page(width=72, height=72)
    with p.open("wb") as f:
        w.write(f)

    doc = read_document(p)
    assert isinstance(doc.text, str)


def test_read_folder_mixed(tmp_path: Path):
    (tmp_path / "a.txt").write_text("alpha beta gamma", encoding="utf-8")

    import docx  # type: ignore

    d = docx.Document()
    d.add_paragraph("alpha beta delta")
    d.save(tmp_path / "b.docx")

    docs = read_folder(tmp_path)
    names = sorted(x.name for x in docs)
    assert "a.txt" in names
    assert "b.docx" in names
