from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional


@dataclass(frozen=True)
class Document:
    name: str
    text: str
    path: Path


SUPPORTED_EXTS = (".txt", ".pdf", ".docx")


def read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as e:  # pragma: no cover
        raise RuntimeError("PDF support requires 'pypdf'. Install it via requirements.txt.") from e

    reader = PdfReader(str(path))
    parts: List[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts)


def read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except Exception as e:  # pragma: no cover
        raise RuntimeError(
            "DOCX support requires 'python-docx'. Install it via requirements.txt."
        ) from e

    d = docx.Document(str(path))
    return "\n".join(p.text for p in d.paragraphs)


def read_document(path: Path) -> Document:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".txt":
        text = read_txt(path)
    elif suffix == ".pdf":
        text = read_pdf(path)
    elif suffix == ".docx":
        text = read_docx(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix} ({path.name})")

    return Document(name=path.name, text=text, path=path)


def read_folder(
    folder: Path,
    *,
    exts: Optional[Iterable[str]] = None,
    recursive: bool = True,
) -> List[Document]:
    folder = Path(folder)
    if not folder.exists():
        return []

    allowed = tuple(e.lower() for e in (exts if exts is not None else SUPPORTED_EXTS))

    it = folder.rglob("*") if recursive else folder.glob("*")
    paths = sorted(
        p for p in it if p.is_file() and p.suffix.lower() in allowed and not p.name.startswith(".")
    )

    docs: List[Document] = []
    for p in paths:
        docs.append(read_document(p))
    return docs
